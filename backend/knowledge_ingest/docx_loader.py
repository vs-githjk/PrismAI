"""DOCX loader using python-docx."""

import io

from docx import Document

from .loaders_base import LoadedDoc, LoaderError


def _table_to_markdown(table) -> str:
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
    return "\n".join(rows)


async def load(content: bytes) -> LoadedDoc:
    try:
        doc = Document(io.BytesIO(content))
    except Exception as exc:
        raise LoaderError(f"Unable to open DOCX: {exc}")

    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        md = _table_to_markdown(table)
        if md.strip():
            parts.append("\n" + md + "\n")

    if not parts:
        raise LoaderError("DOCX is empty.")

    return LoadedDoc(text="\n\n".join(parts), page_metadata=[{"page": 1}])
